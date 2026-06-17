import json
from fastapi import APIRouter, Depends, HTTPException, status, File, UploadFile
from sqlalchemy.orm import Session
from typing import List
import datetime

from app.db.session import get_db
from app.db import models, schemas
from app.services.ai_service import ai_service

router = APIRouter()

# --- HEALTH CHECK ---
@router.get("/health")
def health_check(db: Session = Depends(get_db)):
    try:
        # Check database connectivity
        db.execute(models.Base.metadata.tables["resumes"].select().limit(1))
        db_status = "healthy"
    except Exception as e:
        db_status = f"unhealthy: {str(e)}"
        
    return {
        "status": "online",
        "timestamp": datetime.datetime.utcnow().isoformat(),
        "database": db_status,
        "ai_service": "active" if ai_service.enabled else "mocked"
    }

# --- RESUMES API ---
@router.post("/resumes/generate")
def generate_resume(payload: schemas.ResumeGenerateRequest):
    prompt = (
        f"You are an expert ATS-optimized resume builder. Build a fully structured professional resume based on the following user input prompt.\n"
        f"User Prompt: '{payload.prompt}'\n"
        f"Target Role: '{payload.target_role or 'Professional'}'\n\n"
        f"Respond in a strict valid JSON format. Do not include any markdown backticks, markdown code blocks, or extra text. The JSON must exactly match the schema below:\n"
        f"{{\n"
        f"  \"title\": \"target role resume title (e.g. Senior Frontend Developer Resume)\",\n"
        f"  \"summary\": \"professional summary (2-3 sentences, ATS-optimized, high-impact)\",\n"
        f"  \"contact\": {{\n"
        f"    \"name\": \"Full Name (infer or use placeholder like 'John Doe')\",\n"
        f"    \"email\": \"email address (infer or 'johndoe@email.com')\",\n"
        f"    \"phone\": \"phone number (infer or '+1 (555) 019-2834')\",\n"
        f"    \"linkedin\": \"linkedin url (infer or 'linkedin.com/in/johndoe')\",\n"
        f"    \"github\": \"github url (infer or 'github.com/johndoe')\",\n"
        f"    \"website\": \"personal site url (infer or 'johndoe.dev')\"\n"
        f"  }},\n"
        f"  \"experience\": [\n"
        f"    {{\n"
        f"      \"company\": \"Company Name\",\n"
        f"      \"role\": \"Role Title\",\n"
        f"      \"location\": \"City, State\",\n"
        f"      \"start_date\": \"Month Year\",\n"
        f"      \"end_date\": \"Month Year or Present\",\n"
        f"      \"bullets\": [\n"
        f"        \"STAR bullet point 1 (action verb + task + quantified result)\",\n"
        f"        \"STAR bullet point 2\"\n"
        f"      ]\n"
        f"    }}\n"
        f"  ],\n"
        f"  \"education\": [\n"
        f"    {{\n"
        f"      \"institution\": \"University/College Name\",\n"
        f"      \"degree\": \"Degree earned (e.g. B.S. in Computer Science)\",\n"
        f"      \"location\": \"City, State\",\n"
        f"      \"graduation_date\": \"Month Year\"\n"
        f"    }}\n"
        f"  ],\n"
        f"  \"skills\": [\"Skill 1\", \"Skill 2\", \"Skill 3\"],\n"
        f"  \"projects\": [\n"
        f"    {{\n"
        f"      \"name\": \"Project Name\",\n"
        f"      \"description\": \"One sentence project description\",\n"
        f"      \"tech_stack\": \"Comma separated tech stack (e.g. React, Node.js, SQLite)\",\n"
        f"      \"link\": \"Project link (e.g. github.com/johndoe/project)\"\n"
        f"    }}\n"
        f"  ]\n"
        f"}}"
    )

    response_text = ai_service.generate_content(prompt)

    # Clean response text in case Gemini wraps it in markdown backticks
    if response_text.startswith("```"):
        # Remove first line (e.g. ```json) and last line (```)
        lines = response_text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines[-1].startswith("```"):
            lines = lines[:-1]
        response_text = "\n".join(lines).strip()

    if response_text == "API_KEY_MISSING_MOCK_RESPONSE" or response_text.startswith("ERROR:"):
        # Fallback high-quality mock data based on input keywords
        role = payload.target_role or "Software Engineer"
        prompt_lower = payload.prompt.lower()
        
        name = "Taylor Smith"
        email = "taylor.smith@email.com"
        
        if "frontend" in prompt_lower or "react" in prompt_lower or "web" in prompt_lower:
            role = "Senior Frontend Engineer"
            skills = ["React", "TypeScript", "Tailwind CSS", "JavaScript", "HTML5/CSS3", "Next.js", "Jest", "Git"]
            summary = "Dynamic Frontend Engineer with 4+ years of experience building high-performance, responsive web applications. Expert in React ecosystems, performance tuning, and designing polished glassmorphism SaaS layouts."
            experience = [
                {
                    "company": "PixelCraft Studios",
                    "role": "Frontend Engineer",
                    "location": "New York, NY",
                    "start_date": "Mar 2022",
                    "end_date": "Present",
                    "bullets": [
                        "Redesigned the core e-commerce product flow using React and Tailwind CSS, resulting in a 34% increase in conversion rates.",
                        "Optimized bundle size and assets loading, shaving 1.2s off First Contentful Paint (FCP) and improving Lighthouse score from 71 to 96.",
                        "Mentored 2 junior engineers in TypeScript best practices and state management patterns."
                    ]
                },
                {
                    "company": "CodeFlow Inc",
                    "role": "Junior Web Developer",
                    "location": "Boston, MA",
                    "start_date": "Jun 2020",
                    "end_date": "Feb 2022",
                    "bullets": [
                        "Developed and shipped 15+ pixel-perfect landing pages, driving user acquisition up by 18% over 6 months.",
                        "Integrated REST APIs and payment gateways using stripe SDKs, reducing transaction failures by 8%."
                    ]
                }
            ]
            projects = [
                {
                    "name": "Nexus Dashboard",
                    "description": "A dark-themed analytics command center showing real-time client load metrics.",
                    "tech_stack": "React, Vite, Recharts, Tailwind CSS",
                    "link": "github.com/taylorsmith/nexus-dashboard"
                }
            ]
        elif "backend" in prompt_lower or "python" in prompt_lower or "database" in prompt_lower:
            role = "Backend Systems Engineer"
            skills = ["Python", "FastAPI", "PostgreSQL", "SQLite", "Docker", "Redis", "REST APIs", "AWS", "Git"]
            summary = "Results-driven Backend Engineer specializing in building scalable RESTful APIs, database optimizations, and cloud deployments. Proven track record of improving database query latency by 45%."
            experience = [
                {
                    "company": "DataStream Solutions",
                    "role": "Backend Engineer",
                    "location": "Austin, TX",
                    "start_date": "Nov 2021",
                    "end_date": "Present",
                    "bullets": [
                        "Architected and deployed 20+ secure REST endpoints using FastAPI, supporting a peak traffic load of 15,000 requests per minute.",
                        "Refactored complex PostgreSQL query loops, reducing database response times by 42% and hardware overhead by $2,000 monthly.",
                        "Containerized backend microservices using Docker, decreasing deployment times by 22% in CI/CD pipelines."
                    ]
                }
            ]
            projects = [
                {
                    "name": "TaskFlow API",
                    "description": "An asynchronous task queue manager database API.",
                    "tech_stack": "FastAPI, PostgreSQL, Celery, Redis",
                    "link": "github.com/taylorsmith/taskflow-api"
                }
            ]
        else:
            skills = ["Project Management", "Agile Methodologies", "SQL", "Team Leadership", "Data Analysis", "Python"]
            summary = f"Dedicated and professional candidate targeting a role as a {role}. Experienced in collaborating with cross-functional teams, driving operational efficiencies, and managing deliverables successfully."
            experience = [
                {
                    "company": "Global Innovations",
                    "role": role,
                    "location": "Chicago, IL",
                    "start_date": "Sep 2022",
                    "end_date": "Present",
                    "bullets": [
                        "Successfully directed key deliverables, increasing project turnaround speed by 15% and satisfying SLA targets.",
                        "Leveraged data analytics reports to streamline operations, saving the department approximately 8 hours of manual audit work weekly."
                    ]
                }
            ]
            projects = [
                {
                    "name": "Portfolio Tracker",
                    "description": "A spreadsheet tracker and dashboard reporting operational KPIs.",
                    "tech_stack": "SQL, Excel, Python",
                    "link": "github.com/taylorsmith/portfolio-tracker"
                }
            ]
            
        return {
            "title": f"{role} Resume",
            "summary": summary,
            "contact": {
                "name": name,
                "email": email,
                "phone": "+1 (555) 019-2834",
                "linkedin": "linkedin.com/in/taylorsmith",
                "github": "github.com/taylorsmith",
                "website": "taylorsmith.dev"
            },
            "experience": experience,
            "education": [
                {
                    "institution": "State University",
                    "degree": "B.S. in Computer Science & Information Systems",
                    "location": "Boston, MA",
                    "graduation_date": "May 2020"
                }
            ],
            "skills": skills,
            "projects": projects
        }

    try:
        data = json.loads(response_text)
        return data
    except Exception:
        # If it failed to parse, return a structured fallback with the raw response in summary
        return {
            "title": f"{payload.target_role or 'AI Generated'} Resume",
            "summary": "Completed profile text parsing.",
            "contact": {
                "name": "Jane Doe",
                "email": "janedoe@email.com",
                "phone": "+1 (555) 019-2834",
                "linkedin": "linkedin.com/in/janedoe",
                "github": "github.com/janedoe",
                "website": "janedoe.dev"
            },
            "experience": [
                {
                    "company": "Enterprise Corp",
                    "role": payload.target_role or "Specialist",
                    "location": "Remote",
                    "start_date": "Jan 2023",
                    "end_date": "Present",
                    "bullets": [response_text.strip()[:200]]
                }
            ],
            "education": [
                {
                    "institution": "University of Tech",
                    "degree": "B.S. in Tech Studies",
                    "location": "Online",
                    "graduation_date": "Dec 2022"
                }
            ],
            "skills": ["Communication", "Problem Solving", "Technology"],
            "projects": []
        }

@router.post("/resumes", response_model=schemas.ResumeResponse, status_code=status.HTTP_201_CREATED)
def create_resume(resume: schemas.ResumeCreate, db: Session = Depends(get_db)):
    db_resume = models.Resume(
        title=resume.title,
        summary=resume.summary,
        raw_text=resume.raw_text,
        json_content=resume.json_content
    )
    db.add(db_resume)
    db.commit()
    db.refresh(db_resume)
    return db_resume

@router.get("/resumes", response_model=List[schemas.ResumeResponse])
def get_resumes(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    return db.query(models.Resume).offset(skip).limit(limit).all()

@router.get("/resumes/{resume_id}", response_model=schemas.ResumeResponse)
def get_resume(resume_id: int, db: Session = Depends(get_db)):
    db_resume = db.query(models.Resume).filter(models.Resume.id == resume_id).first()
    if not db_resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return db_resume

@router.delete("/resumes/{resume_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_resume(resume_id: int, db: Session = Depends(get_db)):
    db_resume = db.query(models.Resume).filter(models.Resume.id == resume_id).first()
    if not db_resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    db.delete(db_resume)
    db.commit()
    return None


# --- ENHANCER API ---
@router.post("/enhancements/optimize", response_model=schemas.EnhancementResponse)
def optimize_bullet_point(payload: schemas.EnhancementRequest):
    prompt = (
        f"You are a professional resume writer. Rewrite the following resume bullet point to make it more "
        f"impactful, action-oriented, and quantified using the STAR method (Situation, Task, Action, Result).\n"
        f"Target Job: {payload.job_title or 'Professional'}\n"
        f"Bullet point: '{payload.bullet_point}'\n"
        f"Respond in a strict valid JSON format with keys: 'enhanced' (the optimized bullet point string) and "
        f"'changes_made' (a brief explanation of changes). Do not include any markdown wrapper or backticks."
    )
    
    response_text = ai_service.generate_content(prompt)
    
    if response_text == "API_KEY_MISSING_MOCK_RESPONSE" or response_text.startswith("ERROR:"):
        # Simulated premium output for mock mode
        return schemas.EnhancementResponse(
            original=payload.bullet_point,
            enhanced=f"Successfully spearheaded a cross-functional project, boosting delivery efficiency by 27% and saving 15+ hours weekly through workflow automation.",
            changes_made="Converted passive verbs to strong action verbs, added a clear quantified metric (27% efficiency), and highlighted the business value (saving 15+ hours weekly)."
        )
        
    try:
        data = json.loads(response_text)
        return schemas.EnhancementResponse(
            original=payload.bullet_point,
            enhanced=data.get("enhanced", payload.bullet_point),
            changes_made=data.get("changes_made", "Refactored vocabulary and structure.")
        )
    except Exception:
        # Fallback in case Gemini returns unstructured text
        return schemas.EnhancementResponse(
            original=payload.bullet_point,
            enhanced=response_text.strip(),
            changes_made="Optimized grammar, sentence structure, and vocabulary keywords."
        )


# --- ATS SCAN API ---
@router.post("/ats-scan/analyze", response_model=schemas.AtsScanResponse)
def analyze_ats(scan: schemas.AtsScanCreate, db: Session = Depends(get_db)):
    # Validate resume exists
    db_resume = db.query(models.Resume).filter(models.Resume.id == scan.resume_id).first()
    if not db_resume:
        raise HTTPException(status_code=404, detail="Associated resume not found")

    resume_text = db_resume.raw_text or str(db_resume.json_content)
    
    prompt = (
        f"You are an ATS compliance auditor. Analyze the following candidate profile against the job title: '{scan.job_title or 'General'}' "
        f"and description: '{scan.job_description or 'General'}'\n\n"
        f"Profile Content:\n{resume_text}\n\n"
        f"Rate the compatibility on a scale of 0 to 100.\n"
        f"Respond in strict valid JSON format with keys:\n"
        f"1. 'score': integer between 0 and 100\n"
        f"2. 'keyword_gaps': list of strings representing missing skills/keywords\n"
        f"3. 'formatting_issues': list of strings representing design/formatting flaws\n"
        f"4. 'suggestions': list of strings detailing actionable advice.\n"
        f"Do not include any markdown backticks or extra text."
    )
    
    response_text = ai_service.generate_content(prompt)
    
    if response_text == "API_KEY_MISSING_MOCK_RESPONSE" or response_text.startswith("ERROR:"):
        # Mock ATS Report
        report = {
            "keyword_gaps": ["Kubernetes", "CI/CD Pipelines", "System Design"],
            "formatting_issues": ["Double-column layouts can confuse parser models", "Avoid icons in the contact header"],
            "suggestions": ["Include 3 more bullet points emphasizing cloud infrastructure deployment", "Simplify the typography layout"]
        }
        score = 78
    else:
        try:
            data = json.loads(response_text)
            score = data.get("score", 70)
            report = {
                "keyword_gaps": data.get("keyword_gaps", []),
                "formatting_issues": data.get("formatting_issues", []),
                "suggestions": data.get("suggestions", [])
            }
        except Exception:
            score = 75
            report = {"error": "Failed to parse structured audit feedback", "raw_gemini": response_text}

    db_scan = models.AtsScan(
        resume_id=scan.resume_id,
        job_title=scan.job_title,
        job_description=scan.job_description,
        score=score,
        analysis_report=report
    )
    db.add(db_scan)
    db.commit()
    db.refresh(db_scan)
    return db_scan

@router.get("/ats-scan/history/{resume_id}", response_model=List[schemas.AtsScanResponse])
def get_ats_scans(resume_id: int, db: Session = Depends(get_db)):
    return db.query(models.AtsScan).filter(models.AtsScan.resume_id == resume_id).order_by(models.AtsScan.created_at.desc()).all()


# --- PORTFOLIOS API ---
@router.get("/portfolios/check-slug/{slug}")
def check_slug_availability(slug: str, db: Session = Depends(get_db)):
    cleaned_slug = slug.lower().strip()
    if not cleaned_slug:
        return {"available": False, "suggestion": ""}
    
    # Check if slug is unique
    existing = db.query(models.Portfolio).filter(models.Portfolio.slug == cleaned_slug).first()
    if not existing:
        return {"available": True, "suggestion": ""}
    
    # Generate alternative suggestion
    counter = 1
    while True:
        candidate = f"{cleaned_slug}-{counter}"
        alt_existing = db.query(models.Portfolio).filter(models.Portfolio.slug == candidate).first()
        if not alt_existing:
            return {"available": False, "suggestion": candidate}
        counter += 1

@router.post("/portfolios", response_model=schemas.PortfolioResponse, status_code=status.HTTP_201_CREATED)
def create_portfolio(portfolio: schemas.PortfolioCreate, db: Session = Depends(get_db)):
    requested_slug = portfolio.slug.lower().strip()
    resolved_slug = requested_slug
    
    # Auto-resolve duplicate slug conflicts silently
    existing = db.query(models.Portfolio).filter(models.Portfolio.slug == requested_slug).first()
    if existing:
        counter = 1
        while True:
            candidate = f"{requested_slug}-{counter}"
            alt_existing = db.query(models.Portfolio).filter(models.Portfolio.slug == candidate).first()
            if not alt_existing:
                resolved_slug = candidate
                break
            counter += 1
        
    db_portfolio = models.Portfolio(
        resume_id=portfolio.resume_id,
        slug=resolved_slug,
        title=portfolio.title,
        theme=portfolio.theme,
        customizations=portfolio.customizations
    )
    db.add(db_portfolio)
    db.commit()
    db.refresh(db_portfolio)
    return db_portfolio

@router.get("/portfolios/{slug}", response_model=schemas.PortfolioResponse)
def get_portfolio(slug: str, db: Session = Depends(get_db)):
    db_portfolio = db.query(models.Portfolio).filter(models.Portfolio.slug == slug.lower()).first()
    if not db_portfolio:
        raise HTTPException(status_code=404, detail="Portfolio not found")
    return db_portfolio


# --- COMPARE RESUME VS JD API ---
@router.post("/match-jd/compare", response_model=schemas.JdMatchResponse)
def compare_resume_vs_jd(payload: schemas.JdMatchRequest):
    prompt = (
        f"Compare this candidate's resume content with the job description below:\n"
        f"Resume:\n{payload.resume_text}\n"
        f"Job Description:\n{payload.job_description}\n\n"
        f"Perform matching analysis. Respond in strict valid JSON format with keys:\n"
        f"1. 'match_percentage': integer (0 to 100)\n"
        f"2. 'matched_skills': list of strings\n"
        f"3. 'missing_skills': list of strings\n"
        f"4. 'recommendations': list of strings\n"
        f"Do not include markdown wrappers."
    )
    
    response_text = ai_service.generate_content(prompt)
    
    if response_text == "API_KEY_MISSING_MOCK_RESPONSE" or response_text.startswith("ERROR:"):
        return schemas.JdMatchResponse(
            match_percentage=82,
            matched_skills=["React", "TypeScript", "Tailwind CSS", "Git"],
            missing_skills=["Framer Motion", "GraphQL", "End-to-End Testing (Cypress)"],
            recommendations=[
                "Add details regarding Framer Motion micro-interactions under experience projects.",
                "Mention past exposure to REST APIs as well as GraphQL query integrations."
            ]
        )
        
    try:
        data = json.loads(response_text)
        return schemas.JdMatchResponse(
            match_percentage=data.get("match_percentage", 60),
            matched_skills=data.get("matched_skills", []),
            missing_skills=data.get("missing_skills", []),
            recommendations=data.get("recommendations", [])
        )
    except Exception:
        return schemas.JdMatchResponse(
            match_percentage=70,
            matched_skills=["General Technical Skills"],
            missing_skills=["Specific Job Description Skills"],
            recommendations=["Verify prompt execution logs. Tailor your resume wording to align directly with requirements."]
        )


# --- FEATURE 1: RESUME FILE UPLOAD & PARSING ---
# --- FEATURE 1: RESUME FILE UPLOAD & PARSING ---
@router.post("/resumes/upload")
async def upload_resume_file(file: UploadFile = File(...)):
    filename = file.filename or "resume.pdf"
    content = await file.read()
    file_size_bytes = len(content)
    
    # 1. File Size Validation (Max 10MB)
    max_size = 10 * 1024 * 1024  # 10MB
    if file_size_bytes > max_size:
        raise HTTPException(
            status_code=400, 
            detail=f"File exceeds maximum size of 10MB. Uploaded size: {file_size_bytes / (1024 * 1024):.2f}MB."
        )

    ext = filename.split(".")[-1].lower()
    text_extraction_engine = "text-fallback"
    extracted_text = ""
    ocr_triggered = False
    ocr_log = "N/A"
    
    # 2. Text Extraction Stage
    try:
        if ext == "pdf":
            try:
                import fitz
                # Use PyMuPDF
                doc = fitz.open(stream=content, filetype="pdf")
                text_lines = []
                for page in doc:
                    text_lines.append(page.get_text())
                extracted_text = "\n".join(text_lines).strip()
                text_extraction_engine = "pymupdf"
            except (ImportError, Exception):
                # Fallback to pypdf (which is installed in requirements.txt)
                from pypdf import PdfReader
                import io
                reader = PdfReader(io.BytesIO(content))
                text_lines = []
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_lines.append(page_text)
                extracted_text = "\n".join(text_lines).strip()
                text_extraction_engine = "pypdf"
            
            # Check for OCR fallback if text is scanned/empty
            if len(extracted_text) < 100:
                ocr_triggered = True
                try:
                    from pdf2image import convert_from_bytes
                    import pytesseract
                    
                    images = convert_from_bytes(content)
                    ocr_lines = []
                    for image in images:
                        page_text = pytesseract.image_to_string(image)
                        if page_text:
                            ocr_lines.append(page_text)
                    ocr_text = "\n".join(ocr_lines).strip()
                    
                    if len(ocr_text) >= 100:
                        extracted_text = ocr_text
                        text_extraction_engine = "pytesseract-ocr"
                        ocr_log = "Successfully executed OCR on scanned PDF pages."
                    else:
                        ocr_log = f"OCR completed but output text is too short ({len(ocr_text)} chars)."
                except Exception as ocr_err:
                    ocr_log = f"Scanned document detected, but OCR failed or system binaries (Tesseract/Poppler) missing: {str(ocr_err)}"
        elif ext in ["docx", "doc"]:
            try:
                import docx
                import io
                doc_file = docx.Document(io.BytesIO(content))
                text_lines = []
                for para in doc_file.paragraphs:
                    if para.text:
                        text_lines.append(para.text)
                for table in doc_file.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            text_lines.append(" | ".join(row_text))
                extracted_text = "\n".join(text_lines).strip()
                text_extraction_engine = "python-docx"
            except (ImportError, Exception):
                # Fallback to docx2txt (installed in requirements.txt)
                import docx2txt
                import io
                extracted_text = docx2txt.process(io.BytesIO(content)).strip()
                text_extraction_engine = "docx2txt"
        else:
            # Fallback to plain text decode
            try:
                extracted_text = content.decode("utf-8")
            except Exception:
                extracted_text = content.decode("latin-1")
            extracted_text = extracted_text.strip()
            text_extraction_engine = "utf-decode"
    except Exception as e:
        extracted_text = ""
        ocr_log = f"Text extraction failed: {str(e)}"

    if not extracted_text or len(extracted_text.strip()) < 10:
        raise HTTPException(
            status_code=400,
            detail=f"Unable to extract text from this document. {ocr_log if ocr_triggered else ''}"
        )

    # 3. AI Parsing Stage
    prompt = (
        f"You are a professional resume parser. Parse the following resume text and extract all details into a structured JSON object.\n"
        f"Resume Text:\n{extracted_text}\n\n"
        f"Respond in a strict valid JSON format. Do not include any markdown code blocks (like ```json), no wrappers, no explanations, no text before or after the JSON. "
        f"The JSON must exactly match the schema below:\n"
        f"{{\n"
        f"  \"name\": \"Full Name\",\n"
        f"  \"email\": \"email address\",\n"
        f"  \"phone\": \"phone number\",\n"
        f"  \"location\": \"address or city, state\",\n"
        f"  \"linkedin\": \"linkedin profile url\",\n"
        f"  \"github\": \"github profile url\",\n"
        f"  \"summary\": \"professional resume summary (2-3 sentences)\",\n"
        f"  \"education\": [\n"
        f"    {{\n"
        f"      \"institution\": \"University/College Name\",\n"
        f"      \"degree\": \"Degree earned\",\n"
        f"      \"location\": \"City, State\",\n"
        f"      \"graduation_date\": \"Graduation Year (e.g. 2024)\",\n"
        f"      \"cgpa\": \"CGPA or GPA (if present, e.g. 3.8/4.0)\"\n"
        f"    }}\n"
        f"  ],\n"
        f"  \"skills\": [\"Skill 1\", \"Skill 2\"],\n"
        f"  \"experience\": [\n"
        f"    {{\n"
        f"      \"company\": \"Company Name\",\n"
        f"      \"role\": \"Role Title\",\n"
        f"      \"location\": \"City, State\",\n"
        f"      \"start_date\": \"Month Year\",\n"
        f"      \"end_date\": \"Month Year or Present\",\n"
        f"      \"bullets\": [\n"
        f"        \"accomplishment bullet 1\",\n"
        f"        \"accomplishment bullet 2\"\n"
        f"      ]\n"
        f"    }}\n"
        f"  ],\n"
        f"  \"projects\": [\n"
        f"    {{\n"
        f"      \"name\": \"Project Name\",\n"
        f"      \"description\": \"Project description\",\n"
        f"      \"tech_stack\": \"Tech stack used\",\n"
        f"      \"link\": \"Project link\"\n"
        f"    }}\n"
        f"  ],\n"
        f"  \"certifications\": [\n"
        f"    {{\n"
        f"      \"name\": \"Certification Name\",\n"
        f"      \"issuer\": \"Issuing Authority\"\n"
        f"    }}\n"
        f"  ]\n"
        f"}}"
    )

    ai_parsing_status = "success"
    json_validation_status = "success"
    
    response_text = ai_service.generate_content(prompt)
    raw_ai_response = response_text
    
    # Clean response text in case Gemini wraps it in markdown backticks
    if response_text.startswith("```"):
        lines = response_text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines[-1].startswith("```"):
            lines = lines[:-1]
        response_text = "\n".join(lines).strip()

    parsed_json = {}
    
    # 4. JSON Validation & Fallback Processing
    if response_text == "API_KEY_MISSING_MOCK_RESPONSE" or response_text.startswith("ERROR:"):
        ai_parsing_status = "mocked"
        text_lower = extracted_text.lower()
        role = "Software Engineer"
        if "tally" in text_lower or "account" in text_lower:
            role = "Tally Accountant"
        elif "data" in text_lower or "analyst" in text_lower:
            role = "Data Analyst"
        elif "frontend" in text_lower or "react" in text_lower:
            role = "Senior Frontend Engineer"

        parsed_json = {
            "name": "Alex Mercer",
            "email": "alexmercer@email.com",
            "phone": "+1 (555) 019-2834",
            "location": "San Francisco, CA",
            "linkedin": "linkedin.com/in/alexmercer",
            "github": "github.com/alexmercer",
            "summary": f"Detail-oriented {role} with a proven track record of optimizing workflows, executing data audits, and delivering high-quality client results.",
            "experience": [
                {
                    "company": "Enterprise Tech Solutions",
                    "role": role,
                    "location": "San Francisco, CA",
                    "start_date": "Jan 2022",
                    "end_date": "Present",
                    "bullets": [
                        f"Spearheaded critical processes for the {role} department, boosting productivity by 24% over 6 months.",
                        "Optimized database ledgers and client tracking metrics, saving 15+ hours weekly through python automation scripting."
                    ]
                }
            ],
            "education": [
                {
                    "institution": "State University",
                    "degree": "B.S. in Computer Science & Information Systems",
                    "location": "Boston, MA",
                    "graduation_date": "May 2021",
                    "cgpa": "3.8/4.0"
                }
            ],
            "skills": [
                "React" if "react" in role.lower() else "Tally Prime" if "tally" in role.lower() else "SQL",
                "Python", "MS Excel", "Problem Solving", "Git"
            ],
            "projects": [
                {
                    "name": "Automated Compliance Auditing System",
                    "description": "Interactive analytics dashboard tracking operations KPI delivery logs.",
                    "tech_stack": "Python, Excel, SQLite",
                    "link": "github.com/alexmercer/auditing"
                }
            ],
            "certifications": [
                {
                    "name": "AWS Certified Solutions Architect",
                    "issuer": "Amazon Web Services"
                }
            ]
        }
    else:
        try:
            parsed_json = json.loads(response_text)
        except Exception:
            json_validation_status = "repaired-fallback"
            try:
                cleaned = response_text.strip()
                if cleaned.startswith("[") and cleaned.endswith("]"):
                    cleaned = cleaned[1:-1].strip()
                parsed_json = json.loads(cleaned)
            except Exception:
                parsed_json = {
                    "name": "Alex Mercer",
                    "email": "alexmercer@email.com",
                    "phone": "+1 (555) 019-2834",
                    "location": "San Francisco, CA",
                    "linkedin": "linkedin.com/in/alexmercer",
                    "github": "github.com/alexmercer",
                    "summary": "Completed profile text parsing.",
                    "education": [],
                    "skills": ["Communication", "Problem Solving", "Tech Systems"],
                    "experience": [],
                    "projects": [],
                    "certifications": []
                }

    # Enforce keys
    required_keys = ["name", "email", "phone", "location", "linkedin", "github", "summary", "education", "skills", "experience", "projects", "certifications"]
    for key in required_keys:
        if key not in parsed_json:
            parsed_json[key] = "" if key in ["name", "email", "phone", "location", "linkedin", "github", "summary"] else []

    # Map arrays
    if not isinstance(parsed_json["education"], list): parsed_json["education"] = []
    if not isinstance(parsed_json["skills"], list): parsed_json["skills"] = []
    if not isinstance(parsed_json["experience"], list): parsed_json["experience"] = []
    if not isinstance(parsed_json["projects"], list): parsed_json["projects"] = []
    if not isinstance(parsed_json["certifications"], list): parsed_json["certifications"] = []

    # 5. Quality Score Heuristics
    confidence_score = 50
    if len(extracted_text) > 300: confidence_score += 10
    if not ocr_triggered:
        confidence_score += 15
    else:
        if "Successfully" in ocr_log: confidence_score += 10

    if parsed_json.get("name"): confidence_score += 5
    if parsed_json.get("email"): confidence_score += 5
    if parsed_json.get("phone"): confidence_score += 5
    if parsed_json.get("summary"): confidence_score += 5
    if parsed_json["skills"]: confidence_score += 5
    if parsed_json["experience"]: confidence_score += 5
    confidence_score = min(max(confidence_score, 0), 100)

    # 6. Response Construction
    return {
        "file_upload_status": "success",
        "file_type": ext,
        "file_size_bytes": file_size_bytes,
        "text_extraction_engine": text_extraction_engine,
        "extracted_character_count": len(extracted_text),
        "ocr_triggered": ocr_triggered,
        "ocr_log": ocr_log,
        "ai_parsing_status": ai_parsing_status,
        "json_validation_status": json_validation_status,
        "confidence_score": confidence_score,
        "raw_ai_response": raw_ai_response,
        "resume_data": parsed_json
    }


# --- FEATURE 3: AI ASSISTANT SUGGESTIONS ---
@router.post("/ai-assistant/suggest", response_model=schemas.AssistantSuggestResponse)
def get_ai_suggestions(payload: schemas.AssistantSuggestRequest):
    prompt = (
        f"You are a professional career coach. Provide exactly 3 short, high-impact suggestions for the field: '{payload.field_type}' "
        f"for a candidate targeting the role: '{payload.target_role}'.\n"
        f"Current field value: '{payload.current_text or ''}'\n\n"
        f"Respond in a strict valid JSON format with a single key 'suggestions' holding a list of 3 strings. "
        f"Do not include any markdown wrappers or backticks."
    )
    
    response_text = ai_service.generate_content(prompt)
    
    # Clean response text in case Gemini wraps it in markdown backticks
    if response_text.startswith("```"):
        lines = response_text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines[-1].startswith("```"):
            lines = lines[:-1]
        response_text = "\n".join(lines).strip()

    if response_text == "API_KEY_MISSING_MOCK_RESPONSE" or response_text.startswith("ERROR:"):
        role = payload.target_role or "Professional"
        field = payload.field_type
        
        if field == "summary":
            return {
                "suggestions": [
                    f"Detail-oriented {role} with a proven track record of optimizing workflows, reducing costs, and boosting team efficiency.",
                    f"Results-driven {role} specializing in scalable architecture, system debugging, and database optimization.",
                    f"Highly analytical {role} with 3+ years experience delivering data-driven solutions and collaborating with cross-functional teams."
                ]
            }
        elif field == "achievements":
            return {
                "suggestions": [
                    f"Spearheaded critical task automation, boosting overall operational efficiency by 27% and saving 10+ weekly hours.",
                    f"Led a team of developers to deliver a client portal, expanding user engagement rates by 34% in 3 months.",
                    f"Optimized system query latency by 45%, reducing hardware server overhead by $1,200 monthly."
                ]
            }
        elif field == "skills":
            return {
                "suggestions": ["Problem Solving", "Agile Methodologies", "System Architecture"]
            }
        else:
            return {
                "suggestions": [
                    f"GST Invoice Tracker: Automated ledger audits using Python.",
                    f"Metrics Analytics Engine: Coded real-time user activity logs dashboard."
                ]
            }
            
    try:
        data = json.loads(response_text)
        return data
    except Exception:
        return {
            "suggestions": [
                f"Highly skilled {payload.target_role} specializing in automation, delivery, and scalability.",
                f"Led engineering tasks, saving 8+ developer hours weekly.",
                f"Key skills: System Design, Agile Frameworks, Communication."
            ]
        }


# --- FEATURE 5: AUTO SKILL RECOMMENDATIONS ---
@router.post("/ai-assistant/recommend-skills", response_model=schemas.SkillsRecommendResponse)
def get_skills_recommendation(payload: schemas.RoleRecommendRequest):
    prompt = (
        f"Recommend exactly 8 key skills (a mix of technical and soft skills) for a candidate targeting the role: '{payload.target_role}'.\n"
        f"Respond in a strict valid JSON format with a single key 'skills' containing a list of strings. Do not include markdown backticks."
    )
    
    response_text = ai_service.generate_content(prompt)
    
    if response_text.startswith("```"):
        lines = response_text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines[-1].startswith("```"):
            lines = lines[:-1]
        response_text = "\n".join(lines).strip()

    if response_text == "API_KEY_MISSING_MOCK_RESPONSE" or response_text.startswith("ERROR:"):
        role_lower = payload.target_role.lower()
        if "tally" in role_lower or "account" in role_lower:
            return {"skills": ["Tally Prime", "GST Filing", "TDS Calculation", "MS Excel", "Financial Auditing", "Bank Reconciliation", "Ledger Management", "Problem Solving"]}
        elif "data" in role_lower or "analyst" in role_lower:
            return {"skills": ["SQL", "MS Excel", "Power BI", "Python", "Pandas", "Statistics", "Data Visualization", "Communication"]}
        elif "frontend" in role_lower or "react" in role_lower or "developer" in role_lower:
            return {"skills": ["React", "TypeScript", "Tailwind CSS", "JavaScript", "HTML5/CSS3", "Next.js", "Redux", "Git"]}
        else:
            return {"skills": ["System Analysis", "Technical Writing", "Problem Solving", "Team Collaboration", "Agile Methodologies", "Git", "Project Management", "Data Tracking"]}
            
    try:
        data = json.loads(response_text)
        return data
    except Exception:
        return {"skills": ["SQL", "Python", "Project Management", "Agile Methodologies", "Communication", "Problem Solving"]}


# --- FEATURE 6: AUTO PROJECT RECOMMENDATIONS ---
@router.post("/ai-assistant/recommend-projects", response_model=schemas.ProjectsRecommendResponse)
def get_projects_recommendation(payload: schemas.RoleRecommendRequest):
    prompt = (
        f"Provide exactly 3 custom project ideas tailored for a candidate targeting the role: '{payload.target_role}'.\n"
        f"Respond in a strict valid JSON format with a single key 'projects' holding a list of objects. "
        f"Each object must have exactly keys: 'name', 'description', and 'tech_stack' (comma-separated string). "
        f"Do not include markdown wrappers."
    )
    
    response_text = ai_service.generate_content(prompt)
    
    if response_text.startswith("```"):
        lines = response_text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines[-1].startswith("```"):
            lines = lines[:-1]
        response_text = "\n".join(lines).strip()

    if response_text == "API_KEY_MISSING_MOCK_RESPONSE" or response_text.startswith("ERROR:"):
        role_lower = payload.target_role.lower()
        if "tally" in role_lower or "account" in role_lower:
            return {
                "projects": [
                    {"name": "GST Management System", "description": "Automated GST filing tracker and ledgers auditing tools.", "tech_stack": "Tally Prime, Excel"},
                    {"name": "Inventory Tracking Dashboard", "description": "Designed real-time stock balance sheets and alert logs.", "tech_stack": "Excel, Python"},
                    {"name": "Financial Reporting Automation", "description": "Automated balance sheet compilers reducing human audit error.", "tech_stack": "Python, SQLite"}
                ]
            }
        elif "data" in role_lower or "analyst" in role_lower:
            return {
                "projects": [
                    {"name": "Sales Performance Dashboard", "description": "Interactive business performance dashboard reporting sales KPIs.", "tech_stack": "Power BI, SQL, Python"},
                    {"name": "Customer Churn Analysis Model", "description": "Statistical model identifying high-risk subscriber segments.", "tech_stack": "Python, Pandas, Scikit-learn"},
                    {"name": "HR Analytics Command Center", "description": "Monitored employee churn rates, performance logs, and salary stats.", "tech_stack": "Tableau, Excel"}
                ]
            }
        else:
            return {
                "projects": [
                    {"name": "Metrics Analytics Engine", "description": "Analytics board tracking operations KPI delivery logs.", "tech_stack": "React, Tailwind CSS, Node.js"},
                    {"name": "Task Auditing Queue", "description": "Asynchronous job scheduler matching workloads to worker pools.", "tech_stack": "Python, Celery, Redis"},
                    {"name": "Personal Branding Website", "description": "Responsive multi-theme visual portfolio page compiling achievements.", "tech_stack": "React, Vite, Framer Motion"}
                ]
            }

    try:
        data = json.loads(response_text)
        return data
    except Exception:
        return {
            "projects": [
                {"name": "Custom Command Center", "description": "Analytics board tracking operations KPI delivery logs.", "tech_stack": "React, Node.js"}
            ]
        }


# --- FEATURE 8: ASSISTANT CERTIFICATIONS ---
@router.post("/ai-assistant/recommend-certifications", response_model=schemas.CertificationsRecommendResponse)
def get_certifications_recommendation(payload: schemas.RoleRecommendRequest):
    prompt = (
        f"Recommend 3 professional certifications for a candidate targeting the role: '{payload.target_role}'.\n"
        f"Respond in a strict valid JSON format with a single key 'certifications' containing a list of objects. "
        f"Each object must have exactly keys: 'name' and 'issuer'. Do not include markdown wrappers."
    )
    
    response_text = ai_service.generate_content(prompt)
    
    if response_text.startswith("```"):
        lines = response_text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines[-1].startswith("```"):
            lines = lines[:-1]
        response_text = "\n".join(lines).strip()

    if response_text == "API_KEY_MISSING_MOCK_RESPONSE" or response_text.startswith("ERROR:"):
        role_lower = payload.target_role.lower()
        if "tally" in role_lower or "account" in role_lower:
            return {
                "certifications": [
                    {"name": "Certified Tally Professional", "issuer": "Tally Solutions"},
                    {"name": "Chartered Financial Analyst (Level 1)", "issuer": "CFA Institute"},
                    {"name": "QuickBooks Certified User", "issuer": "Certiport"}
                ]
            }
        else:
            return {
                "certifications": [
                    {"name": "AWS Certified Solutions Architect", "issuer": "Amazon Web Services"},
                    {"name": "Google Cloud Professional Data Engineer", "issuer": "Google Cloud"},
                    {"name": "Project Management Professional (PMP)", "issuer": "Project Management Institute"}
                ]
            }


# --- FEATURE: AI COVER LETTER GENERATOR ---
@router.post("/cover-letters/generate")
def generate_cover_letter(payload: schemas.CoverLetterGenerateRequest, db: Session = Depends(get_db)):
    resume_data = {}
    if payload.resume_id:
        db_resume = db.query(models.Resume).filter(models.Resume.id == payload.resume_id).first()
        if db_resume:
            resume_data = db_resume.json_content
    elif payload.resume_data:
        resume_data = payload.resume_data

    name = resume_data.get("contact", {}).get("name", "") if isinstance(resume_data.get("contact"), dict) else resume_data.get("name", "John Doe")
    role = resume_data.get("title", "Professional")
    if "Resume" in role:
        role = role.replace(" Resume", "")
        
    job_description = payload.job_description or ""

    prompt = (
        f"You are a professional cover letter writer. Create a highly customized cover letter based on the resume data and target job description.\n"
        f"Resume Data: {json.dumps(resume_data)}\n"
        f"Job Description: {job_description or 'N/A'}\n"
        f"Letter Type: {payload.letter_type}\n"
        f"Writing Style: {payload.style}\n\n"
        f"Respond in a strict valid JSON format. Do not include markdown wrappers. The JSON must exactly match the schema below:\n"
        f"{{\n"
        f"  \"title\": \"Cover Letter Title (e.g. {role} Cover Letter)\",\n"
        f"  \"greeting\": \"Greeting (e.g. Dear Hiring Team,)\",\n"
        f"  \"introduction\": \"An engaging introduction paragraph matching the {payload.style} tone, explaining interest in the role.\",\n"
        f"  \"body\": \"A strong body section highlighting matching experience, skills, and projects aligned with the Job Description. Use key words from the job description.\",\n"
        f"  \"closing\": \"A professional closing call-to-action paragraph.\",\n"
        f"  \"signature\": \"Sincerely,\\n{name}\",\n"
        f"  \"score\": 85,\n"
        f"  \"personalization_score\": 90,\n"
        f"  \"ats_score\": 80,\n"
        f"  \"tone_score\": 85,\n"
        f"  \"structure_score\": 90,\n"
        f"  \"keywords_detected\": [\"keyword1\", \"keyword2\"]\n"
        f"}}"
    )

    response_text = ai_service.generate_content(prompt)

    # Clean response text in case Gemini wraps it in markdown backticks
    if response_text.startswith("```"):
        lines = response_text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines[-1].startswith("```"):
            lines = lines[:-1]
        response_text = "\n".join(lines).strip()

    if response_text == "API_KEY_MISSING_MOCK_RESPONSE" or response_text.startswith("ERROR:"):
        detected_keywords = ["React", "TypeScript", "FastAPI", "Ledgers", "GST", "SQL"]
        matched_kws = [kw for kw in detected_keywords if kw.lower() in job_description.lower()] if job_description else ["Analytical", "Agile", "Teamwork"]
        
        intro = f"I am writing to express my strong interest in the open position as detailed in your job posting. With my extensive background in the industry, I am confident in my ability to make a significant contribution to your organization."
        body_text = f"Throughout my career, I have honed my skills and successfully delivered on complex projects. My background aligns well with the requirements, and I am eager to apply my experience in React development, system designs, and team collaborations to help drive success at your company."
        closing_text = f"Thank you for considering my application. I look forward to the opportunity to discuss how my background and qualifications align with your organization's needs."
        
        if "tally" in role.lower() or "account" in role.lower():
            intro = f"I am excited to apply for the Accountant position. With my solid experience in bookkeeping, ledgers auditing, and financial reporting, I am prepared to bring strong accounting accuracy to your team."
            body_text = f"My experience includes managing ledgers, executing bank reconciliations, and ensuring GST compliance. I have successfully minimized errors by utilizing advanced Excel tools, contributing to overall financial efficiency."
        elif "data" in role.lower() or "analyst" in role.lower():
            intro = f"I am writing to apply for the Data Analyst role. I have a passion for transforming complex raw metrics into clear, actionable business insights to drive strategic decisions."
            body_text = f"In my past work, I have developed interactive SQL databases, compiled KPI reports, and automated tracking dashboards. This experience has allowed me to identify optimization workflows and save teams valuable audit hours weekly."

        return {
            "title": f"{role} Cover Letter",
            "letter_type": payload.letter_type,
            "style": payload.style,
            "content": {
                "title": f"{role} Cover Letter",
                "greeting": f"Dear Hiring Manager,",
                "introduction": intro,
                "body": body_text,
                "closing": closing_text,
                "signature": f"Best Regards,\n{name}"
            },
            "score": 88,
            "personalization_score": 85,
            "ats_score": 90,
            "tone_score": 87,
            "structure_score": 90,
            "keywords_detected": matched_kws
        }

    try:
        data = json.loads(response_text)
        content_keys = ["title", "greeting", "introduction", "body", "closing", "signature"]
        content_dict = {}
        for k in content_keys:
            content_dict[k] = data.get(k, "")
            
        return {
            "title": data.get("title", f"{role} Cover Letter"),
            "letter_type": payload.letter_type,
            "style": payload.style,
            "content": content_dict,
            "score": data.get("score", 85),
            "personalization_score": data.get("personalization_score", 85),
            "ats_score": data.get("ats_score", 80),
            "tone_score": data.get("tone_score", 85),
            "structure_score": data.get("structure_score", 90),
            "keywords_detected": data.get("keywords_detected", [])
        }
    except Exception:
        return {
            "title": f"{role} Cover Letter",
            "letter_type": payload.letter_type,
            "style": payload.style,
            "content": {
                "title": f"{role} Cover Letter",
                "greeting": "Dear Hiring Team,",
                "introduction": f"I am pleased to submit my application for the target role. My experience and skill set align perfectly with your organization's goals.",
                "body": f"I possess strong analytical and problem-solving skills, with a track record of successful project deliveries. I am confident in my ability to add value to your team.",
                "closing": "Thank you for your time and consideration.",
                "signature": f"Sincerely,\n{name}"
            },
            "score": 75,
            "personalization_score": 70,
            "ats_score": 75,
            "tone_score": 80,
            "structure_score": 80,
            "keywords_detected": []
        }

@router.post("/cover-letters/optimize", response_model=schemas.CoverLetterSection)
def optimize_cover_letter(payload: schemas.CoverLetterOptimizeRequest):
    prompt = (
        f"You are an expert ATS optimizer. Optimize the following cover letter content to increase matches for the job description.\n"
        f"Job Description: {payload.job_description}\n"
        f"Cover Letter Greeting: {payload.current_content.greeting}\n"
        f"Introduction: {payload.current_content.introduction}\n"
        f"Body: {payload.current_content.body}\n"
        f"Closing: {payload.current_content.closing}\n"
        f"Signature: {payload.current_content.signature}\n\n"
        f"Respond in a strict valid JSON format with keys: 'greeting', 'introduction', 'body', 'closing', 'signature'. "
        f"Do not include markdown wrappers."
    )
    response_text = ai_service.generate_content(prompt)
    if response_text.startswith("```"):
        lines = response_text.splitlines()
        if lines[0].startswith("```"):
            lines = lines[1:]
        if lines[-1].startswith("```"):
            lines = lines[:-1]
        response_text = "\n".join(lines).strip()

    if response_text == "API_KEY_MISSING_MOCK_RESPONSE" or response_text.startswith("ERROR:"):
        return payload.current_content

    try:
        data = json.loads(response_text)
        return {
            "title": payload.current_content.title,
            "greeting": data.get("greeting", payload.current_content.greeting),
            "introduction": data.get("introduction", payload.current_content.introduction),
            "body": data.get("body", payload.current_content.body),
            "closing": data.get("closing", payload.current_content.closing),
            "signature": data.get("signature", payload.current_content.signature)
        }
    except Exception:
        return payload.current_content

@router.post("/cover-letters/improve")
def improve_cover_letter_section(payload: schemas.CoverLetterImproveRequest):
    prompt = (
        f"You are a professional editor. Improve the following cover letter text according to this instruction: '{payload.instruction}'.\n"
        f"Original Text: {payload.text}\n\n"
        f"Return ONLY the improved text, no explanations, no headers."
    )
    response_text = ai_service.generate_content(prompt)
    if response_text == "API_KEY_MISSING_MOCK_RESPONSE" or response_text.startswith("ERROR:"):
        return {"improved_text": payload.text + " (AI Optimized: Aligned to style objectives)"}
    return {"improved_text": response_text.strip()}

# CRUD routes
@router.post("/cover-letters", response_model=schemas.CoverLetterResponse, status_code=status.HTTP_201_CREATED)
def save_cover_letter(cover_letter: schemas.CoverLetterCreate, db: Session = Depends(get_db)):
    db_letter = models.CoverLetter(
        resume_id=cover_letter.resume_id,
        title=cover_letter.title,
        letter_type=cover_letter.letter_type,
        style=cover_letter.style,
        content=cover_letter.content.dict(),
        score=cover_letter.score,
        personalization_score=cover_letter.personalization_score,
        ats_score=cover_letter.ats_score,
        tone_score=cover_letter.tone_score,
        structure_score=cover_letter.structure_score,
        keywords_detected=cover_letter.keywords_detected
    )
    db.add(db_letter)
    db.commit()
    db.refresh(db_letter)
    return db_letter

@router.get("/cover-letters", response_model=List[schemas.CoverLetterResponse])
def get_cover_letters(db: Session = Depends(get_db)):
    return db.query(models.CoverLetter).order_by(models.CoverLetter.created_at.desc()).all()

@router.delete("/cover-letters/{letter_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_cover_letter(letter_id: int, db: Session = Depends(get_db)):
    db_letter = db.query(models.CoverLetter).filter(models.CoverLetter.id == letter_id).first()
    if not db_letter:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    db.delete(db_letter)
    db.commit()
    return

